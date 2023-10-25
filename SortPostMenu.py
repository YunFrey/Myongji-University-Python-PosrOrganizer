# Pyside6 로드
from PySide6.QtWidgets import *
from PySide6.QtCore import *
from PySide6.QtGui import *

# ui 로드용
from PySide6 import QtUiTools
from ui_loader import load_ui
# SQLite 로드
import sqlite3
# 파일 관리자
import os
# 시간 관련 라이브러리
from datetime import datetime
import time
# 판다스 라이브러리
import pandas as pd
# python-docx 라이브러리
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class SortPostMenu(QMainWindow):
    def __init__(self):
        super().__init__()
        load_ui('SortPostMenu.ui', self)
        self.show()

        #SQLite 연결
        if os.path.isfile("postlist.db"):  #DB파일이 있을 경우
            self.listdb = sqlite3.connect("postlist.db")
            self.cursor = self.listdb.cursor()
        else:   #DB파일이 없을 경우
            self.statusBar.showMessage('DB가 없으므로 DB파일을 생성합니다.')
            self.listdb = sqlite3.connect("postlist.db")
            self.cursor = self.listdb.cursor()
            self.cursor.execute(
                "CREATE TABLE postlist(id integer PRIMARY KEY AUTOINCREMENT, 접수날짜 text, 사원번호 text, 접수자명 text, 부서명 text, 보내는사람 text, 받는사람 text, 주소 text, 우편번호 text, 제목 text, 수량 integer, 종류 text, 긴급여부 boolean, 가로길이 real, 세로길이 real, 우편중량 real, 전화번호 text, 높이 real, 메모 text, 할인타입 text, 할인타입그룹 integer, 결재여부 , 반려여부 boolean)")
            self.listdb.commit()

        #QprogressBar 숨겨놓기(군집화정렬때만 표시)
        self.progressBar.hide()

        # 버튼 시그널 함수로 연결(제작자정보메뉴)
        self.action_closewindow.triggered.connect(self.closewindow)

        # 접수된 우편물 조회
        self.btn_loadpostlist.clicked.connect(self.loadpostlist)

        self.datesel_start.setDate(QDate.currentDate())
        self.datesel_end.setDate(QDate.currentDate())
        self.datesel_start.setCalendarPopup(True)
        self.datesel_end.setCalendarPopup(True)
        # [QDateWidget 설정 2] -시작날짜가 끝 날짜와 겹치지 않게 날짜 변경시에 보정
        self.datesel_start.dateChanged.connect(lambda: self.leftdatefix())
        self.datesel_end.dateChanged.connect(lambda: self.rightdatefix())

        # 전체 선택 버튼
        self.btn_selectallrow_up.clicked.connect(self.selectallrowup)
        # 전체 선택 버튼
        self.btn_selectallrow_down.clicked.connect(self.selectallrowdown)

        #결재대상 버튼
        self.btn_setpostready.clicked.connect(self.setpostready)
        #반려버튼
        self.btn_returnpost.clicked.connect(self.returnpostselected)
        #군집화정렬버튼
        self.btn_startkmeans.clicked.connect(self.start_kmeans_sort)
        #규격정렬버튼
        self.btn_organizepost.clicked.connect(self.organizepost)
        #리스트 출력 버튼
        self.btn_makedocx.clicked.connect(self.df_to_docx)
        #결재대상에서 우편물 삭제버튼
        self.btn_remove_selpost.clicked.connect(self.remove_selectedpost)
        #결재버튼
        self.btn_postend.clicked.connect(self.fee_to_docx)

        #요금계산 버튼
        self.btn_calculatefee.clicked.connect(self.addpostfee)
        # 이 창을 열었을 때 DB 조회 및 생성

        #묶음할인버튼 시그널
        self.btn_grouppost.clicked.connect(self.grouppost)

        # 메인윈도우 X 눌러 종료 시 DB 정리
    def closeEvent(self, event):
        self.listdb.close()
        self.close()
    def closewindow(self): #윈도우 닫기
        self.listdb.close()
        self.close()

    def rightdatefix(self):
        if(int(self.datesel_start.date().toString('yyyyMMdd'))) > (int(self.datesel_end.date().toString('yyyyMMdd'))):
            self.datesel_start.setDate(self.datesel_end.date())
    def leftdatefix(self):
        if (int(self.datesel_end.date().toString('yyyyMMdd'))) < (int(self.datesel_start.date().toString('yyyyMMdd'))):
            self.datesel_end.setDate(self.datesel_start.date())

    def loadpostlist(self):
        #SQL 로드
        if os.path.isfile("postlist.db"):  # DB파일이 있을 경우
            listdb = sqlite3.connect("postlist.db")
            cursor = listdb.cursor()
        else:  # DB파일이 없을 경우
            self.statusBar().showMessage('DB가 없으므로 DB파일을 생성합니다.')
            listdb = sqlite3.connect("postlist.db")
            cursor = listdb.cursor()
            cursor.execute(
                "CREATE TABLE postlist(id integer PRIMARY KEY AUTOINCREMENT, 접수날짜 text, 사원번호 text, 접수자명 text, 부서명 text, 보내는사람 text, 받는사람 text, 주소 text, 우편번호 text, 제목 text, 수량 integer, 종류 text, 긴급여부 text, 가로길이 real, 세로길이 real, 우편중량 real, 전화번호 text, 높이 real, 메모 text, 할인타입 text, 할인타입그룹 integer, 결재여부 , 반려여부 boolean, 비용 text)")
            listdb.commit()

        #시작날짜 변수에 담기
        str_datestart = self.datesel_start.date().toString("yyyyMMdd")
        print(str_datestart)
        #끝나는날짜 변수에 남기
        str_dateend = self.datesel_end.date().toString("yyyyMMdd")
        print(str_dateend)
        # SQL DB 를 QTableWidget 에 표시
        #오늘날짜 저장
        todaydat: str = datetime.today().strftime("%Y%m%d")
        # 결재된 우편물도 조회하는 옵션 켜져잇을 경우
        if self.info_14_isprocessed.isChecked() == True:
            optional = "AND 결재여부 = '1'"
        else:
            optional = "AND 결재여부 IS NULL"
        #쿼리 완성
        query = 'SELECT id, 접수날짜, 사원번호, 접수자명, 부서명, 보내는사람, 받는사람, 주소, 우편번호, 제목, 수량, 종류, 긴급여부, 가로길이, 세로길이, 우편중량, 전화번호, 높이, 메모, 할인타입, 할인타입그룹, 결재여부, 반려여부, 비용  FROM postlist WHERE "접수날짜" >= %s AND "접수날짜" <= %s %s' % (str_datestart, str_dateend, optional)
        try:
            df = pd.read_sql(query, listdb)
        except pd.errors.DatabaseError:
            print('시스템 : 접수된 우편물이 없습니다.')
        except NameError:
            print('시스템 : 접수된 우편물이 없습니다.')
        else:
            print('[쿼리]\n', query)

            if df.empty == True:  # 데이터프레임이 비어있을 떄
                # 오류 메세지 표시
                self.table_postlist.clear()  # 초기화
                self.statusBar().showMessage('알림 : 조회된 건이 없습니다.')
            else:
                # QTableWidget 에 DF 내용 쓰기
                self.table_postlist.clear()  # 부르기 전 초기화
                col = len(df.keys())
                self.table_postlist.setColumnCount(col)
                self.table_postlist.setHorizontalHeaderLabels(df.keys())
                row = len(df.index)
                self.table_postlist.setRowCount(row)
                for r in range(row):
                    for c in range(col):
                        item = QTableWidgetItem(str(df.iloc[r][c]))
                        self.table_postlist.setItem(r, c, item)
                self.table_postlist.resizeColumnsToContents()


    # 선택된 줄에 대한 모든 데이터를 읽는 함수
    def getselectedrowid_adv(self, tablename):
        #인텍스에 선택된 롤 집어넣기
        index = tablename.selectionModel().selectedRows()
        postlist = list()
        for row in index: #index를 정렬하여 row 로 읽기
            # 리스트 초기화
            templist = list()
            for i in range(tablename.columnCount()):
                # 가로로 값 긁기
                tempvar = tablename.item(row.row(), i).text()
                # 수집된 우편물 정보 list 에 append
                templist.append(tempvar)
            postlist.append(templist) #우편물 정보 append
        return postlist

    # 선택된 줄에 대한 모든 데이터를 읽는 함수(1레코드만 읽음)
    def getselectedrowid_adv_one(self, tablename):
        # 인텍스에 선택된 롤 집어넣기
        index = tablename.selectionModel().selectedRows()
        # 리스트 초기화
        templist = list()
        for row in index:  # index를 정렬하여 row 로 읽기
            # 읽힌 모든 거 다 append 하기
            for i in range(tablename.columnCount()):
                tempvar = tablename.item(row.row(), i).text()
                templist.append(tempvar)
        return templist

        #선택된 우편물 반려해서 다른 DB에 집어넣기
    def returnpostselected(self):
        # 반려 사유 묻기
        try:
            if len(self.table_postlist.selectionModel().selectedRows()) != 1:
                raise Exception
        except Exception as e:  # 예외가 발생했을 때 실행됨
            msg_noentry = QMessageBox()  # 메세지박스 생성
            msg_noentry.warning(self, "오류", "반려할 우편물을 1개만 선택해 주세요")  # 메세지박스 설정
            output_noentry = msg_noentry.show()  # 메세지박스 실행
            if output_noentry == QMessageBox.Ok:  # OK 누르면 닫기
                msg_noentry.close()
        else :
            # 창 클래스 구성
            class msg_rtncausewin(QDialog):
                def __init__(self):
                    super().__init__()
                    # SQLite 불러오기
                    if os.path.isfile("postlist.db"):  # DB파일이 있을 경우
                        self.listdb = sqlite3.connect("postlist.db")
                        self.cursor = self.listdb.cursor()
                    else:  # DB파일이 없을 경우
                        self.statusBar.showMessage('DB가 없으므로 DB파일을 생성합니다.')
                        self.listdb = sqlite3.connect("postlist.db")
                        self.cursor = self.listdb.cursor()
                        self.cursor.execute(
                            "CREATE TABLE postlist(id integer PRIMARY KEY AUTOINCREMENT, 접수날짜 text, 사원번호 text, 접수자명 text, 부서명 text, 보내는사람 text, 받는사람 text, 주소 text, 우편번호 text, 제목 text, 수량 integer, 종류 text, 긴급여부 boolean, 가로길이 real, 세로길이 real, 우편중량 real, 전화번호 text, 높이 real, 메모 text, 할인타입 text, 할인타입그룹 integer, 결재여부 , 반려여부 boolean)")
                        self.listdb.commit()

                    # 윈도우 설정
                    self.setWindowIcon(QIcon('title_big.png'))
                    self.setWindowTitle('반려사유')
                    self.resize(300, 100)
                    # Label 추가
                    self.desc_label = QLabel('반려사유를 입력하세요')
                    # 중앙정렬
                    self.desc_label.setAlignment(Qt.AlignCenter)
                    # LineEdit 추가
                    self.entrybox = QLineEdit()
                    self.entrybox.move(115, 30)
                    # 버튼 추가
                    self.btn_closewin = QPushButton("OK")
                    self.btn_closewin.resize(150, 50)

                    # 레이아웃 생성 및 연결
                    self.layout = QVBoxLayout()
                    self.layout.addWidget(self.desc_label)
                    self.layout.addWidget(self.entrybox)
                    self.layout.addWidget(self.btn_closewin)
                    self.setLayout(self.layout)

                    # 버튼 시그널 연결
                    self.btn_closewin.clicked.connect(self.set_rtncause)

                    # 창 생성 및 루프
                    self.exec()

                def set_rtncause(self):
                    # entrybox 내용 불러와서 클래스변수로 저장
                    SortPostMenu.rtncause = str(self.entrybox.text())
                    self.close()

            # [클래스 종료 시 수행]
            # 객체 생성
            self.window = msg_rtncausewin()
            # 반려우편물 DB 로드
            if os.path.isfile("returnpostlist.db"):  # DB파일이 있을 경우
                return_listdb = sqlite3.connect("returnpostlist.db")
                return_cursor = return_listdb.cursor()
            else:  # DB파일이 없을 경우
                self.statusBar().showMessage('반려 DB가 없으므로 DB파일을 생성합니다.')
                return_listdb = sqlite3.connect("returnpostlist.db")
                return_cursor = return_listdb.cursor()
                return_cursor.execute(
                    "CREATE TABLE postlist(id integer, 접수날짜 text, 사원번호 text, 접수자명 text, 부서명 text, 보내는사람 text, 받는사람 text, 주소 text, 우편번호 text, 제목 text, 수량 integer, 종류 text, 긴급여부 boolean, 가로길이 real, 세로길이 real, 우편중량 real, 전화번호 text, 높이 real, 메모 text, 할인타입 text, 할인타입그룹 integer, 결재여부 , 반려여부 boolean, 비용 text)")
                return_listdb.commit()

            # SQL INSERT 문 실행
            postline = self.getselectedrowid_adv_one(self.table_postlist)
            postline[18] = SortPostMenu.rtncause
            print(postline)
            try:
                return_cursor.execute(
                    "INSERT INTO postlist(id, 접수날짜, 사원번호, 접수자명, 부서명, 보내는사람, 받는사람, 주소, 우편번호, 제목, 수량, 종류, 긴급여부, 가로길이, 세로길이, 우편중량, 전화번호, 높이, 메모, 할인타입, 할인타입그룹, 결재여부, 반려여부, 비용) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                    postline)

            except sqlite3.ProgrammingError:
                msg_forbidrtn = QMessageBox()  # 메세지박스 생성
                msg_forbidrtn.warning(self, "오류", '여러개를 한번에 반려할 수는 없습니다.')  # 메세지박스 경고옵션 설정
                output_forbidrtn = msg_forbidrtn.show()  # 메세지박스 실행
                if output_forbidrtn == QMessageBox.Ok:  # OK 누르면 닫기
                    msg_forbidrtn.close()
                self.statusBar().showMessage('여러개를 한번에 반려할 수는 없습니다. ')

            return_listdb.commit()

            # SQL DELETE 문 실행
            idlist = self.getselectedrowid(self.table_postlist)
            self.cursor.execute("DELETE FROM postlist WHERE id IN (%s)" % ",".join(idlist))
            # 커밋
            self.listdb.commit()
            # 테이블 업데이트
            self.table_postlist_sel.clear()
            self.loadpostlist()



        #선택된 줄에 대한 ID 데이터를 읽는 함수
    def getselectedrowid(self, tablename):
        #인텍스에 선택된 롤 집어넣기
        index = tablename.selectionModel().selectedRows()
        #리스트 초기화
        templist = list()
        for row in index: #index를 정렬하여 row 로 읽기
            #row 정보 중 0번칼럼 읽어서 텍스트로 출력
            id = tablename.item(row.row(), 0).text()
            #선택된 모든 row 읽어서 리스트에 추가
            templist.append(id)
        return templist

    #결제대상 우편물로 넘기는 함수
    def setpostready(self):
        # 선택한 우편이 있는지 검사
        try:
            if len(self.table_postlist.selectionModel().selectedRows()) == 0:
                raise Exception
        except Exception as e:  # 예외가 발생했을 때 실행됨
            msg_noentry = QMessageBox()  # 메세지박스 생성
            msg_noentry.warning(self, "오류", "결재할 우편물을 선택해 주세요")  # 메세지박스 설정
            output_noentry = msg_noentry.show()  # 메세지박스 실행
            if output_noentry == QMessageBox.Ok:  # OK 누르면 닫기
                msg_noentry.close()
                #선택된 우편이 있을 경우
        else:
            # 선택한 우편물의 정보 기록 긁어오기
            readypostlist = self.getselectedrowid_adv(self.table_postlist)
            # 결재테이블 내 모든 데이터 선택
            self.table_postlist.selectAll()
            # 데이터프레임에 헤더값 넣기
            labels = []
            for c in range(self.table_postlist.columnCount()):
                headeritem = self.table_postlist.horizontalHeaderItem(c).text()
                labels.append(headeritem)
            # 데이터프레임 생성
            checkdf = pd.DataFrame(readypostlist)
            checkdf.columns = labels

            # 결재대상 우편 있을 경우 플래그 설정
            if '1' in checkdf['결재여부'].values:
                isforbid = True
            else:
                isforbid = False

            if isforbid == True:
                # 오류 메세지박스 생성(단순객체) #
                msg_forbid = QMessageBox()  # 메세지박스 생성
                msg_forbid.warning(self, "오류", '이미 결재된 우편물은 결재대상이 될 수 없습니다.')  # 메세지박스 경고옵션 설정
                output_forbid = msg_forbid.show()  # 메세지박스 실행
                if output_forbid == QMessageBox.Ok:  # OK 누르면 닫기
                    msg_forbid.close()
            else:
                # 추후 묶음할인기능이 수행된 상태인지 여부를 저장하는 변수(아직 사용안했기에 False)
                self.is_groupified = False

                # 데이터프레임에 헤더값 넣기
                labels = []
                for c in range(self.table_postlist.columnCount()):
                    headeritem = self.table_postlist.horizontalHeaderItem(c).text()
                    labels.append(headeritem)

                # 데이터프레임 생성
                postdf = pd.DataFrame(readypostlist)
                try:
                    postdf.columns = labels
                except ValueError as e:
                    # 우편물을 선택하지 않아 postdf의 라벨과 오차가 발생할 때 오류핸들링
                    self.statusBar().showMessage('에러 : 우편물을 선택해 주세요')

                print('[데이터프레임]\n', postdf)

                # 우편물 중량을 25,50,100,150 등으로 정렬

                # QTableWidget 에 DF 내용 쓰기
                self.table_postlist_sel.clear()  # 부르기 전 초기화
                col = len(postdf.keys())
                self.table_postlist_sel.setColumnCount(col)
                self.table_postlist_sel.setHorizontalHeaderLabels(postdf.keys())
                row = len(postdf.index)
                self.table_postlist_sel.setRowCount(row)
                for r in range(row):
                    for c in range(col):
                        item = QTableWidgetItem(str(postdf.iloc[r][c]))
                        self.table_postlist_sel.setItem(r, c, item)
                self.table_postlist_sel.resizeColumnsToContents()




    def start_kmeans_sort(self):
        #군집화 모델은 우편물 여러개가 있어야 가능하기에 우편물 최소 5개를 결재할때만 사용가능하게함
        print('갯수:', self.table_postlist_sel.rowCount())
        if(self.table_postlist_sel.rowCount() < 5):
            self.statusBar().showMessage('에러 : 군집화정렬 기능은 우편물 5개 이상일때만 사용가능합니다.')
        else:
            # 결재테이블 내 모든 데이터 선택
            self.table_postlist_sel.selectAll()
            # 결재대상 데이터 긁어오기
            readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
            # 데이터프레임에 헤더값 넣기
            labels = []
            for c in range(self.table_postlist_sel.columnCount()):
                headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
                labels.append(headeritem)
            # 데이터프레임 생성
            postdf = pd.DataFrame(readypostlist)
            postdf.columns = labels
            print('[데이터프레임_kmeansinput]\n', postdf)
            #원핫인코딩
            #오리지널 칼럼 백업
            postdf2 = pd.get_dummies(postdf, columns=['종류'])
            print('[데이터프레임_kmeansinput_2]\n', postdf2.to_string())
            # Kmeans,실루엣 분석 라이브러리 호출
            from sklearn.cluster import KMeans
            from sklearn.metrics import silhouette_score
            # 이상치 있을 때 오류 안나야하므로 스케일러 호출(효과 확인불가로 미사용)
            #from sklearn.preprocessing import StandardScaler
            #scaler = StandardScaler()
            #postdf = scaler.fit_transform()

            # 경고문 제거
            import warnings
            from sklearn.exceptions import DataConversionWarning
            warnings.filterwarnings(action='ignore', category=DataConversionWarning)
            from sklearn.exceptions import ConvergenceWarning
            warnings.filterwarnings(action='ignore', category=ConvergenceWarning)

            # Input 값 가공
            rangelist = list(range(12,15)) + list(range(17,17)) + list(range(23, len(postdf2.columns)))

            X = postdf2.iloc[:, rangelist]
            print('[K-mean input data]\n', X.describe())

            #QprogressBar 리셋
            self.progressBar.reset()
            self.progressBar.show()
            self.statusBar().clearMessage()

            def setprogress_percentage(input):
                percentage = (input/(X.shape[0]-2))*100
                self.progressBar.setValue(percentage)

            # 실루엣 분석을 통한 최적 군집수 값 탐색
            silhouette_scores = []
            # 군집화 기능은 보내는 우편물이 5개 이상일 때만 가능
            print('xshape',X.shape[0])
            for k in range(4, X.shape[0]-1): #분류는 5개부터 우편물 개수개 까지 탐색
                setprogress_percentage(k)
                print('Kmean 클러스터 개수별 실루엣점수 스캔 : ',k)
                kmeanModel = KMeans(n_init=10, n_clusters=k, random_state=12345).fit(X)
                score = silhouette_score(X, kmeanModel.labels_)
                silhouette_scores.append(score)
            print('실루엣 점수 리스트 :', silhouette_scores)
            # 최적 K값 저장
            optimal_k = silhouette_scores.index(max(silhouette_scores)) + 4 #2부터 X.shape[0]까지 탐색했으므로
            print('최적 k값 : ', optimal_k)

            # kmeans 모델 생성 및 데이터 적합
            kmeans = KMeans(n_init=10, n_clusters=optimal_k, random_state=12345)
            kmeans.fit(X)

            print('후',postdf2)

            # 분류데이터 오리지널 postdf에 추가
            postdf['분류그룹'] = kmeans.labels_
            # 알고리즘에서 분류한 값대로 정렬 후 우편중량으로 세부정렬
            postdf = postdf.sort_values(['분류그룹', '종류', '우편중량'])

            # QTableWidget 에 DF 내용 쓰기
            self.table_postlist_sel.clear()  # 부르기 전 초기화
            col = len(postdf.keys())
            self.table_postlist_sel.setColumnCount(col)
            self.table_postlist_sel.setHorizontalHeaderLabels(postdf.keys())
            row = len(postdf.index)
            self.table_postlist_sel.setRowCount(row)
            for r in range(row):
                for c in range(col):
                    item = QTableWidgetItem(str(postdf.iloc[r][c]))
                    self.table_postlist_sel.setItem(r, c, item)
            self.table_postlist_sel.resizeColumnsToContents()

    def fixweight(self, input):
        if float(input) <= 5:
            res = 5
        elif float(input) <= 25:
            res = 25
        elif float(input) <= 50:
            res = 50
        elif float(input) <= 1000:
            res = 1000
        elif float(input) <= 2000:
            res = 2000
        elif float(input) <= 6000:
            res = 6000
        else:
            res = input
        return res

    def fixweight_package(self, input):
        if float(input) <= 3000:
            res = 3000
        elif float(input) <= 5000:
            res = 5000
        elif float(input) <= 7000:
            res = 7000
        elif float(input) <= 10000:
            res = 10000
        elif float(input) <= 15000:
            res = 15000
        elif float(input) <= 20000:
            res = 20000
        else:
            res = input
        return res

    def standardsorter(self):
        pass

    #우편물 중량,가로세로,높이 정보에 맞춰 우편물을 규격,규격외로 분류
    def organizepost(self):

        # 결재테이블 내 모든 데이터 선택
        self.table_postlist_sel.selectAll()
        print('결재대상 우편물 건수 :', len(self.getselectedrowid(self.table_postlist_sel)))
        if (len(self.getselectedrowid(self.table_postlist_sel))) != 0:
            # 결재대상 데이터 긁어오기
            readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
            # 데이터프레임에 헤더값 넣기
            labels = []
            for c in range(self.table_postlist_sel.columnCount()):
                headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
                labels.append(headeritem)
            # 데이터프레임 생성
            postdf = pd.DataFrame(readypostlist)
            postdf.columns = labels
            print(postdf.to_string())

            ###########################################
            # [1차 정렬] : 중량 규격 정리하기(25,50,100...)
            postdf.loc[(postdf['종류'] == '우편'), '우편중량'] = postdf.우편중량.map(self.fixweight)
            postdf.loc[(postdf['종류'] == '등기'), '우편중량'] = postdf.우편중량.map(self.fixweight)
            postdf.loc[(postdf['종류'] == '익일특급'), '우편중량'] = postdf.우편중량.map(self.fixweight)
            #postdf["우편중량"] = postdf.우편중량.map(self.fixweight)
            # [1차 소포정렬]
            postdf.loc[(postdf['종류'] == '일반소포'), '우편중량'] = postdf.우편중량.map(self.fixweight_package)
            postdf.loc[(postdf['종류'] == '등기소포'), '우편중량'] = postdf.우편중량.map(self.fixweight_package)
            ###########################################

            ###########################################
            # [2차 정렬 : 가로세로높이에 따른 규격우편과 규격외 우편 분류]
            # 1차적으로 모든 우편물을 규격외로 변경
            postdf.loc[(postdf['종류'] == '우편') , '종류'] = '규격외우편'

            # 규격 우편물 기준에 맞는 우편물을 규격우편으로 분류
            postdf.loc[(postdf['세로길이'].astype(float) >= 9) &
                       (postdf['세로길이'].astype(float) <= 13) &
                       (postdf['가로길이'].astype(float) >= 14) &
                       (postdf['가로길이'].astype(float) <= 23.5) &
                       (postdf['높이'].astype(float) >= 0) &
                       (postdf['높이'].astype(float) <= 0.5) &
                       (postdf['우편중량'].astype(float) <= 50) &
                       (postdf['종류'] == '규격외우편'), '종류'] = '규격우편'
            ############################################

            ############################################
            # [3차 정렬 : 가로세로높이에 따른 규격등기와 규격외 등기 분류]
            # 1차적으로 모든 등기를 규격외등기로 변경
            postdf.loc[(postdf['종류'] == '등기'), '종류'] = '규격외등기'

            # 규격 우편물 기준에 맞는 등기를 규격등기로 분류
            postdf.loc[(postdf['세로길이'].astype(float) >= 9) &
                       (postdf['세로길이'].astype(float) <= 13) &
                       (postdf['가로길이'].astype(float) >= 14) &
                       (postdf['가로길이'].astype(float) <= 23.5) &
                       (postdf['높이'].astype(float) >= 0) &
                       (postdf['높이'].astype(float) <= 0.5) &
                       (postdf['우편중량'].astype(float) <= 50) &
                       (postdf['종류'] == '규격외등기'), '종류'] = '규격등기'
            ############################################

            # [4차 정렬 : 가로세로높이에 따른 규격익일특급과 규격외 익일특급 분류]
            # 1차적으로 모든 등기를 규격외등기로 변경
            postdf.loc[(postdf['종류'] == '익일특급'), '종류'] = '규격외익일특급'

            # 규격 우편물 기준에 맞는 등기를 규격등기로 분류
            postdf.loc[(postdf['세로길이'].astype(float) >= 9) &
                       (postdf['세로길이'].astype(float) <= 13) &
                       (postdf['가로길이'].astype(float) >= 14) &
                       (postdf['가로길이'].astype(float) <= 23.5) &
                       (postdf['높이'].astype(float) >= 0) &
                       (postdf['높이'].astype(float) <= 0.5) &
                       (postdf['우편중량'].astype(float) <= 50) &
                       (postdf['종류'] == '규격외익일특급'), '종류'] = '규격익일특급'
            ############################################


            # QTableWidget 에 DF 내용 쓰기
            self.table_postlist_sel.clear()  # 부르기 전 초기화
            col = len(postdf.keys())
            self.table_postlist_sel.setColumnCount(col)
            self.table_postlist_sel.setHorizontalHeaderLabels(postdf.keys())
            row = len(postdf.index)
            self.table_postlist_sel.setRowCount(row)
            for r in range(row):
                for c in range(col):
                    item = QTableWidgetItem(str(postdf.iloc[r][c]))
                    self.table_postlist_sel.setItem(r, c, item)
            self.table_postlist_sel.resizeColumnsToContents()
        else:
            self.statusBar().showMessage('에러 : 최소 한개의 우편물을 결재대상으로 올려주세요.')

    # 요금 붙이기 실행
    def addpostfee(self):
        #postinformation 데이터가 있는지 확인
        postdb = sqlite3.connect("postinformation.db")
        cursor = postdb.cursor()
        # postage 존재하는지 Execute
        try :
            cursor.execute('SELECT * FROM postage LIMIT 1')
        except :
            #오류 발생 시
            msg_nodb = QMessageBox()
            msg_nodb.warning(self, "오류", "우체국 요금데이터가 없습니다. 갱신해 주세요.")
            output_nodb = msg_nodb.show()
            if output_nodb == QMessageBox.Ok:
                msg_nodb.close()
                self.mainmenu_statusBar.showMessage('리셋완료')
        else :
            #문제 없을 시 실행
            feedb = sqlite3.connect("postlist.db")
            feecursor = feedb.cursor()
            # 결재테이블 내 모든 데이터 선택
            self.table_postlist_sel.selectAll()
            print('요금붙일 우편물 건수 :', len(self.getselectedrowid(self.table_postlist_sel)))
            if (len(self.getselectedrowid(self.table_postlist_sel))) != 0:
                # 데이터 긁어오기
                readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
                # 데이터프레임에 헤더값 넣기
                labels = []
                for c in range(self.table_postlist_sel.columnCount()):
                    headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
                    labels.append(headeritem)
                # 데이터프레임 생성
                print('현재라벨 :', labels)

                postdf = pd.DataFrame(readypostlist)
                postdf.columns = labels
                print('[요금계산대상 데이터프레임 표시]\n', postdf.to_string())

                # 규격우편물 불러와서 df에 저장
                standard_db = sqlite3.connect("postinformation.db")
                standard_cursor = standard_db.cursor()
                query = 'SELECT 보통우편요금 FROM postage'
                standard_df = pd.read_sql(query, standard_db)

                # 문자열에서 필요한 원 앞 숫자만 추출하는 함수
                def extractnumonly(inputstr):
                    result = ""
                    for char in reversed(inputstr):
                        if char.isdigit():
                            result = char + result
                        elif char == "다":
                            break
                    return int(result)

                #### 규격우편물 요금 ####
                # 5g까지
                fee_1 = extractnumonly(standard_df.iloc[0, 0])
                # 5g초과 25g 까지
                fee_2 = extractnumonly(standard_df.iloc[1, 0])
                # 25g초과 50g 까지
                fee_3 = extractnumonly(standard_df.iloc[2, 0])
                #### 규격외우편물 요금 ####
                # 50g 까지
                fee_4 = extractnumonly(standard_df.iloc[3, 0])
                # 50g 초과 1kg 까지
                fee_add_5 = extractnumonly(standard_df.iloc[4, 0])
                # 1kg 초과 2kg 까지
                fee_add_6 = extractnumonly(standard_df.iloc[5, 0])
                # 2kg 초과 6kg 까지
                fee_add_7 = extractnumonly(standard_df.iloc[6, 0])

                ###########################################
                # 규격우편물에 요금 붙이기
                postdf.loc[(postdf['종류'] == '규격우편') &
                           (postdf['우편중량'].astype(float) <= 5), '비용'] = fee_1
                postdf.loc[(postdf['종류'] == '규격우편') &
                           (postdf['우편중량'].astype(float) > 5) &
                           (postdf['우편중량'].astype(float) <= 25), '비용'] = fee_2
                postdf.loc[(postdf['종류'] == '규격우편') &
                           (postdf['우편중량'].astype(float) > 25) &
                           (postdf['우편중량'].astype(float) <= 50), '비용'] = fee_3

                # 규격외우편물에 요금 붙이기

                #### [규격외우편물 중량에 따른 요금계산] ####
                def calc_nonstandard_fee_50(inputfee):
                    addon = (inputfee - 50) // 50
                    print(addon)
                    return (fee_4 + fee_add_5 * addon)

                def calc_nonstandard_fee_1000(inputfee):
                    addon = (inputfee - 1000) // 200
                    return (fee_4 + fee_add_5 * 19 + fee_add_6 * addon)

                def calc_nonstandard_fee_2000(inputfee):
                    addon = (inputfee - 2000) // 1000
                    return (fee_4 + fee_add_5 * 19 + fee_add_6 * 5 + fee_add_7 * addon)

                ###########################################

                ## 추후 .0이 붙는게 싫으면 return 되는 값에다가 str(int())만 붙이면 됨
                postdf.loc[(postdf['종류'] == '규격외우편') &
                           (postdf['우편중량'].astype(float) <= 50), '비용'] = fee_4
                postdf.loc[(postdf['종류'] == '규격외우편') &
                           (postdf['우편중량'].astype(float) > 50) &
                           (postdf['우편중량'].astype(float) <= 1000), '비용'] = calc_nonstandard_fee_50(
                    postdf['우편중량'].astype(float))
                postdf.loc[(postdf['종류'] == '규격외우편') &
                           (postdf['우편중량'].astype(float) > 1000) &
                           (postdf['우편중량'].astype(float) <= 2000), '비용'] = calc_nonstandard_fee_1000(
                    postdf['우편중량'].astype(float))
                postdf.loc[(postdf['종류'] == '규격외우편') &
                           (postdf['우편중량'].astype(float) > 2000) &
                           (postdf['우편중량'].astype(float) <= 6000), '비용'] = calc_nonstandard_fee_2000(
                    postdf['우편중량'].astype(float))
                postdf.loc[(postdf['종류'] == '규격외우편') &
                           (postdf['우편중량'].astype(float) > 6000), '비용'] = '오류'

                # 규격우편물 불러와서 df에 저장
                register_db = sqlite3.connect("postinformation.db")
                register_cursor = register_db.cursor()
                query = 'SELECT 수수료액 FROM postage_special'
                register_df = pd.read_sql(query, register_db)
                print('regggg', register_df.to_string())
                # 수수료액을 가져와 순수 숫자로만 저장하기 (2,100원 > 2100)
                register_fee = register_df.iloc[0, 0]
                # 필터링 적용한다음 char 합치기
                register_fee = int(''.join(filter(str.isdigit, register_fee)))

                # 규격등기에 우편물 요금 붙이기 (규격우편물 비용에 등기수수료만 추가)
                postdf.loc[(postdf['종류'] == '규격등기') &
                           (postdf['우편중량'].astype(float) <= 5), '비용'] = str(fee_1 + register_fee)  # 2100은 등기수수료
                postdf.loc[(postdf['종류'] == '규격등기') &
                           (postdf['우편중량'].astype(float) > 5) &
                           (postdf['우편중량'].astype(float) <= 25), '비용'] = str(fee_2 + register_fee)  # 2100은 등기수수료
                postdf.loc[(postdf['종류'] == '규격등기') &
                           (postdf['우편중량'].astype(float) > 25) &
                           (postdf['우편중량'].astype(float) <= 50), '비용'] = str(fee_3 + register_fee)  # 2100은 등기수수료

                ###########################################

                # 규격외등기에 우편물 요금 붙이기 (규격외우편물 비용에 등기수수료만 추가)
                postdf.loc[(postdf['종류'] == '규격외등기') &
                           (postdf['우편중량'].astype(float) <= 50), '비용'] = fee_4 + register_fee  # 2100은 등기수수료
                postdf.loc[(postdf['종류'] == '규격외등기') &
                           (postdf['우편중량'].astype(float) > 50) &
                           (postdf['우편중량'].astype(float) <= 1000), '비용'] = calc_nonstandard_fee_50(
                    postdf['우편중량'].astype(float)) + register_fee
                postdf.loc[(postdf['종류'] == '규격외등기') &
                           (postdf['우편중량'].astype(float) > 1000) &
                           (postdf['우편중량'].astype(float) <= 2000), '비용'] = calc_nonstandard_fee_1000(
                    postdf['우편중량'].astype(float)) + register_fee
                postdf.loc[(postdf['종류'] == '규격외등기') &
                           (postdf['우편중량'].astype(float) > 2000) &
                           (postdf['우편중량'].astype(float) <= 6000), '비용'] = calc_nonstandard_fee_2000(
                    postdf['우편중량'].astype(float)) + register_fee
                postdf.loc[(postdf['종류'] == '규격외등기') &
                           (postdf['우편중량'].astype(float) > 6000), '비용'] = '오류'

                ###########################################
                # 일반소포 요금 DB에서 가져오기
                package_normal_db = sqlite3.connect("postinformation.db")
                package_normal_cursor = package_normal_db.cursor()
                query = 'SELECT * FROM postage_package_normal'
                package_normal_df = pd.read_sql(query, package_normal_db)

                ##########################################

                # 일반소포 요금 붙이기
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) <= 80) &
                           (postdf['세로길이'].astype(float) <= 80) &
                           (postdf['우편중량'].astype(float) <= 3000), '비용'] = package_normal_df.iloc[0, 2]
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) > 80) &
                           (postdf['가로길이'].astype(float) <= 100) &
                           (postdf['세로길이'].astype(float) > 80) &
                           (postdf['세로길이'].astype(float) <= 100) |
                           (postdf['우편중량'].astype(float) > 3000) &
                           (postdf['우편중량'].astype(float) <= 5000), '비용'] = package_normal_df.iloc[0, 3]
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) > 80) &
                           (postdf['가로길이'].astype(float) <= 100) &
                           (postdf['세로길이'].astype(float) > 80) &
                           (postdf['세로길이'].astype(float) <= 100) |
                           (postdf['우편중량'].astype(float) > 5000) &
                           (postdf['우편중량'].astype(float) <= 7000), '비용'] = package_normal_df.iloc[0, 4]
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) > 100) &
                           (postdf['가로길이'].astype(float) <= 120) &
                           (postdf['세로길이'].astype(float) > 100) &
                           (postdf['세로길이'].astype(float) <= 120) |
                           (postdf['우편중량'].astype(float) > 7000) &
                           (postdf['우편중량'].astype(float) <= 10000), '비용'] = package_normal_df.iloc[0, 5]
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) > 100) &
                           (postdf['가로길이'].astype(float) <= 120) &
                           (postdf['세로길이'].astype(float) > 100) &
                           (postdf['세로길이'].astype(float) <= 120) |
                           (postdf['우편중량'].astype(float) > 10000) &
                           (postdf['우편중량'].astype(float) <= 15000), '비용'] = package_normal_df.iloc[0, 6]
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) > 100) &
                           (postdf['가로길이'].astype(float) <= 120) &
                           (postdf['세로길이'].astype(float) > 100) &
                           (postdf['세로길이'].astype(float) <= 120) |
                           (postdf['우편중량'].astype(float) > 15000) &
                           (postdf['우편중량'].astype(float) <= 20000), '비용'] = package_normal_df.iloc[0, 7]
                postdf.loc[(postdf['종류'] == '일반소포') &
                           (postdf['가로길이'].astype(float) > 120) |
                           (postdf['세로길이'].astype(float) > 120) |
                           (postdf['우편중량'].astype(float) > 15000), '비용'] = '접수불가'

                ###########################################

                # 등기소포 요금 DB에서 가져오기
                package_reg_db = sqlite3.connect("postinformation.db")
                package_reg_cursor = package_reg_db.cursor()
                query = 'SELECT * FROM postage_package_registered'
                package_reg_df = pd.read_sql(query, package_reg_db)

                # 등기소포 요금 붙이기
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) <= 80) &
                           (postdf['세로길이'].astype(float) <= 80) &
                           (postdf['우편중량'].astype(float) <= 3000), '비용'] = package_reg_df.iloc[0, 3]
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) > 80) &
                           (postdf['가로길이'].astype(float) <= 100) &
                           (postdf['세로길이'].astype(float) > 80) &
                           (postdf['세로길이'].astype(float) <= 100) |
                           (postdf['우편중량'].astype(float) > 3000) &
                           (postdf['우편중량'].astype(float) <= 5000), '비용'] = package_reg_df.iloc[0, 4]
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) > 80) &
                           (postdf['가로길이'].astype(float) <= 100) &
                           (postdf['세로길이'].astype(float) > 80) &
                           (postdf['세로길이'].astype(float) <= 100) |
                           (postdf['우편중량'].astype(float) > 5000) &
                           (postdf['우편중량'].astype(float) <= 7000), '비용'] = package_reg_df.iloc[0, 5]
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) > 100) &
                           (postdf['가로길이'].astype(float) <= 120) &
                           (postdf['세로길이'].astype(float) > 100) &
                           (postdf['세로길이'].astype(float) <= 120) |
                           (postdf['우편중량'].astype(float) > 7000) &
                           (postdf['우편중량'].astype(float) <= 10000), '비용'] = package_reg_df.iloc[0, 6]
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) > 100) &
                           (postdf['가로길이'].astype(float) <= 120) &
                           (postdf['세로길이'].astype(float) > 100) &
                           (postdf['세로길이'].astype(float) <= 120) |
                           (postdf['우편중량'].astype(float) > 10000) &
                           (postdf['우편중량'].astype(float) <= 15000), '비용'] = package_reg_df.iloc[0, 7]
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) > 100) &
                           (postdf['가로길이'].astype(float) <= 120) &
                           (postdf['세로길이'].astype(float) > 100) &
                           (postdf['세로길이'].astype(float) <= 120) |
                           (postdf['우편중량'].astype(float) > 15000) &
                           (postdf['우편중량'].astype(float) <= 20000), '비용'] = package_reg_df.iloc[0, 8]
                postdf.loc[(postdf['종류'] == '등기소포') &
                           (postdf['가로길이'].astype(float) > 120) |
                           (postdf['세로길이'].astype(float) > 120) |
                           (postdf['우편중량'].astype(float) > 20000), '비용'] = '접수불가'

                ###########################################
                # 익일특급 수수료액 가져오기
                # 수수료액을 가져와 순수 숫자로만 저장하기 (1,000원 > 1000)
                nextday_fee = register_df.iloc[17, 0]

                # 필터링 적용한다음 char 합치기
                nextday_fee = int(''.join(filter(str.isdigit, nextday_fee)))

                ##########################################

                # 규격익일특급에 우편물 요금 붙이기 (규격우편물 비용에 익일특급수수료만 추가)
                postdf.loc[(postdf['종류'] == '규격익일특급') &
                           (postdf['우편중량'].astype(float) <= 5), '비용'] = fee_1 + nextday_fee  # 1000은 익일특급수수료
                postdf.loc[(postdf['종류'] == '규격익일특급') &
                           (postdf['우편중량'].astype(float) > 5) &
                           (postdf['우편중량'].astype(float) <= 25), '비용'] = fee_2 + nextday_fee  # 1000은 익일특급수수료
                postdf.loc[(postdf['종류'] == '규격익일특급') &
                           (postdf['우편중량'].astype(float) > 25) &
                           (postdf['우편중량'].astype(float) <= 50), '비용'] = fee_3 + nextday_fee  # 1000은 익일특급수수료

                ###########################################

                # 규격외익일특급에 우편물 요금 붙이기 (규격외우편물 비용에 등기수수료만 추가)
                postdf.loc[(postdf['종류'] == '규격외익일특급') &
                           (postdf['우편중량'].astype(float) <= 50), '비용'] = str(fee_4 + nextday_fee)  # 1000은 익일특급수수료
                postdf.loc[(postdf['종류'] == '규격외익일특급') &
                           (postdf['우편중량'].astype(float) > 50) &
                           (postdf['우편중량'].astype(float) <= 1000), '비용'] = calc_nonstandard_fee_50(
                    postdf['우편중량'].astype(float)) + nextday_fee
                postdf.loc[(postdf['종류'] == '규격외익일특급') &
                           (postdf['우편중량'].astype(float) > 1000) &
                           (postdf['우편중량'].astype(float) <= 2000), '비용'] = calc_nonstandard_fee_1000(
                    postdf['우편중량'].astype(float)) + nextday_fee
                postdf.loc[(postdf['종류'] == '규격외익일특급') &
                           (postdf['우편중량'].astype(float) > 2000) &
                           (postdf['우편중량'].astype(float) <= 6000), '비용'] = calc_nonstandard_fee_2000(
                    postdf['우편중량'].astype(float)) + nextday_fee
                postdf.loc[(postdf['종류'] == '규격외익일특급') &
                           (postdf['우편중량'].astype(float) > 6000), '비용'] = '오류'

                ###########################################
                # 마지막 : 사전접수 할인이 가능한 우편물

                #############################################
                # 결과물 콘솔에 표시
                print('[결과 데이터프레임]\n', postdf.to_string())
                ###########################################

                # QTableWidget 에 업데이트된 DF 내용 쓰기
                self.table_postlist_sel.clear()  # 부르기 전 초기화
                col = len(postdf.keys())
                self.table_postlist_sel.setColumnCount(col)
                self.table_postlist_sel.setHorizontalHeaderLabels(postdf.keys())
                row = len(postdf.index)
                self.table_postlist_sel.setRowCount(row)
                for r in range(row):
                    for c in range(col):
                        item = QTableWidgetItem(str(postdf.iloc[r][c]))
                        self.table_postlist_sel.setItem(r, c, item)
                self.table_postlist_sel.resizeColumnsToContents()
            else:
                self.statusBar().showMessage('에러 : 최소 한개의 우편물을 결재대상으로 올려주세요.')
        finally :
            #모든 게 끝나면
            postdb.close() #확인용 변수 정리





    # 전체선택 함수(아래쪽 테이블)
    def selectallrowdown(self):
        self.table_postlist_sel.selectAll()

    #전체선택 함수(위쪽 테이블)
    def selectallrowup(self):
        self.table_postlist.selectAll()

    #그룹화 함수
    def grouppost(self):
        # 데이터 긁어오기
        self.table_postlist_sel.selectAll()
        try:
            if len(self.table_postlist_sel.selectionModel().selectedRows()) == 0:
                raise Exception
        except Exception as e:  # 예외가 발생했을 때 실행됨
            msg_noentry = QMessageBox()  # 메세지박스 생성
            msg_noentry.warning(self, "오류", "묶음할인할 우편물을 선택해 주세요")  # 메세지박스 설정
            output_noentry = msg_noentry.show()  # 메세지박스 실행
            if output_noentry == QMessageBox.Ok:  # OK 누르면 닫기
                msg_noentry.close()
        else:
            #선택된 우편물이 있을 경우
            # 묶음할인 기능이 진행됬음을 저장하는 Flag = True
            self.is_groupified = True
            readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
            print('[res]', readypostlist)
            # 데이터프레임에 헤더값 넣기
            labels = []
            for c in range(self.table_postlist_sel.columnCount()):
                headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
                labels.append(headeritem)
            # 데이터프레임 생성
            print('현재라벨 :', labels)
            postdf = pd.DataFrame(readypostlist)
            postdf.columns = labels

            print('입력 데이터프레임 \n', postdf.to_string())

            # 종류와 우편중량과 요금 칼럼의 값이 같은 레코드들끼리 묶어서 그룹 번호 부여
            postdf['할인타입그룹'] = postdf.groupby(['종류', '우편중량', '비용']).ngroup()

            # 그룹별로 번호를 붙였으니 이제 각 칼럼 유니크값의 빈도수 계산하여 리스트에 담기
            counts = postdf['할인타입그룹'].value_counts()

            # 할인타입그룹에서 동일 값 3개 이상인 레코드는 냅두고 그게 아니면 공백으로 처리
            print('할인타입그룹 정렬시작')
            postdf['할인타입그룹'] = postdf.loc[postdf['종류'] == '규격등기', '할인타입그룹'].apply(
                lambda x: str(x) if counts[x] >= 10 else '')
            # 모든 우편물에 대해 타입 분류하는 코드(미사용)
            # postdf['할인타입그룹'] = postdf['할인타입그룹'].apply(lambda x: x if counts[x] >= 3 else '')

            # QTableWidget 에 DF 내용 쓰기
            self.table_postlist_sel.clear()  # 부르기 전 초기화
            col = len(postdf.keys())
            self.table_postlist_sel.setColumnCount(col)
            self.table_postlist_sel.setHorizontalHeaderLabels(postdf.keys())
            row = len(postdf.index)
            self.table_postlist_sel.setRowCount(row)
            for r in range(row):
                for c in range(col):
                    item = QTableWidgetItem(str(postdf.iloc[r][c]))
                    self.table_postlist_sel.setItem(r, c, item)
            self.table_postlist_sel.resizeColumnsToContents()


    #결재대상의 우편물 삭제
    def remove_selectedpost(self):
        # 선택된 Row의 id값 저장
        indexlist = self.getselectedrowid(self.table_postlist_sel)

        # 데이터 긁어오기
        self.table_postlist_sel.selectAll()
        readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
        print('[res]', readypostlist)
        # 데이터프레임에 헤더값 넣기
        labels = []
        for c in range(self.table_postlist_sel.columnCount()):
            headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
            labels.append(headeritem)
        # 데이터프레임 생성
        print('현재라벨 :', labels)
        postdf = pd.DataFrame(readypostlist)
        postdf.columns = labels

        # 이 기능은 [4. 묶음할인 설정] 기능이 실행된 상태면 실행도지 말아야 함
        if(self.is_groupified == True):
            msg_cantdel = QMessageBox()  # 메세지박스 생성
            msg_cantdel.information(self, "알림", "묶음할인을 진행한 상태에서는 사용 불가합니다. 처음부터 다시 진행해주세요.")  # 메세지박스 설정
            output_cantdel = msg_cantdel.show()  # 메세지박스 실행
            if output_cantdel == QMessageBox.Ok:  # OK 누르면 닫기
                msg_cantdel.close()
                self.mainmenu_statusBar().showMessage('리셋완료')
        else:

            # indexlist 에 있는 레코드들 DF에서 삭제
            # ~연산자로 조건을 반대로 만들어 indexlist 에 없는 우편물만 살리기
            postdf = postdf[~postdf['id'].isin(indexlist)]


            # QTableWidget 에 DF 내용 쓰기
            self.table_postlist_sel.clear()  # 부르기 전 초기화
            col = len(postdf.keys())
            self.table_postlist_sel.setColumnCount(col)
            self.table_postlist_sel.setHorizontalHeaderLabels(postdf.keys())
            row = len(postdf.index)
            self.table_postlist_sel.setRowCount(row)
            for r in range(row):
                for c in range(col):
                    item = QTableWidgetItem(str(postdf.iloc[r][c]))
                    self.table_postlist_sel.setItem(r, c, item)
            self.table_postlist_sel.resizeColumnsToContents()


    #문서화 함수
    def df_to_docx(self):
        # postinformation 데이터가 있는지 확인
        postdb = sqlite3.connect("postinformation.db")
        cursor = postdb.cursor()
        # postage 존재하는지 Execute
        try:
            cursor.execute('SELECT * FROM postage LIMIT 1')
        except:
            # 오류 발생 시
            msg_nodb = QMessageBox()
            msg_nodb.warning(self, "오류", "우체국 요금데이터가 없습니다. 갱신해 주세요.")
            output_nodb = msg_nodb.show()
            if output_nodb == QMessageBox.Ok:
                msg_nodb.close()
                self.mainmenu_statusBar.showMessage('리셋완료')
        else:
            postdb.close()
            # QTableWidget에서 postdf 만들기
            # 데이터 긁어오기
            self.table_postlist_sel.selectAll()
            # 선택된 우편물이 있는지 확인
            try:
                if len(self.table_postlist_sel.selectionModel().selectedRows()) == 0:
                    raise Exception
            except Exception as e:  # 예외가 발생했을 때 실행됨
                msg_noentry = QMessageBox()  # 메세지박스 생성
                msg_noentry.warning(self, "오류", "출력할 우편물이 없습니다.")  # 메세지박스 설정
                output_noentry = msg_noentry.show()  # 메세지박스 실행
                if output_noentry == QMessageBox.Ok:  # OK 누르면 닫기
                    msg_noentry.close()
            else:
                # 우편물이 있을 경우
                readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
                print('[res]', readypostlist)
                # 데이터프레임에 헤더값 넣기
                labels = []
                for c in range(self.table_postlist_sel.columnCount()):
                    headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
                    labels.append(headeritem)
                # 데이터프레임 생성
                print('현재라벨 :', labels)
                postdf = pd.DataFrame(readypostlist)
                postdf.columns = labels

                # 출력할 데이터프레임 검증
                try:  # 데이터프레임 호출
                    postdf = postdf[
                        ['id', '보내는사람', '받는사람', '주소', '우편번호', '제목', '수량', '종류', '우편중량', '전화번호', '할인타입그룹', '비용']]

                except NameError:  # 데이터프레임이 존재하지 않아 못불러올때
                    self.statusBar().showMessage('에러 : 결재대상우편을 지정해주세요.')

                except KeyError:  # 데이터프레임은 있으나 원하는 게 아닐 때
                    self.statusBar().showMessage('에러 : 결재대상우편을 지정해주세요.')

                else:  # 데이터프레임이 존재할때
                    # Word 문서 생성
                    doc = Document()

                    # 제목 추가하기
                    doc.add_heading('결재우편물 리스트', 0)

                    # 페이지 레이아웃 가로로 변경 (미사용)
                    section = doc.sections[-1]
                    doc_width, doc_height = section.page_height, section.page_width
                    section.page_width = doc_width
                    section.page_height = doc_height

                    # 표 생성
                    posttable = doc.add_table(postdf.shape[0] + 1, postdf.shape[1])

                    # 표 테두리 생성
                    posttable.style = 'Table Grid'

                    # 칼럼 추가
                    for i, colname in enumerate(postdf.columns):
                        posttable.cell(0, i).text = colname

                    # 표의 2번쨰 줄부터 레코드 추가
                    for i in range(postdf.shape[0]):
                        for j in range(postdf.shape[1]):
                            posttable.cell(i + 1, j).text = str(postdf.values[i, j])

                    # 테이블의 글자크기를 7로 변경 (코드참조 : https://stackoverflow.com/questions/43007725/python-docx-how-to-change-table-font-size)
                    for row in posttable.rows:  # 테이블의 각 줄
                        for cell in row.cells:  # 테이블 각 줄의 셀
                            paragraphs = cell.paragraphs  # 테이블 각 줄의 셀의 문단그룹
                            for paragraph in paragraphs:  # 테이블 각 줄의 셀의 문단
                                for word in paragraph.runs:  # 문단의 모든 글자
                                    word.font.size = Pt(8)

                    # 시간날짜 마지막에 추가하기
                    text_to_add = "갱신 : " + datetime.today().strftime('%Y년 %m월 %d일 %H시 %M분 %S초')
                    description_last = doc.add_paragraph(text_to_add)
                    description_last.alignment = 2  # 가운데 정렬

                    # Word 문서 저장
                    savename = datetime.today().strftime('%Y-%m-%d_%H-%M-%S_') + '우편물리스트.docx'
                    print('저장된 파일명 :', savename)

                    # 저장 다이얼로그 생성
                    savedialog = QFileDialog.getSaveFileName(self, '우편물 리스트 문서 저장', savename, 'word 파일(*.docx)')

                    # 저장 실행
                    try:
                        doc.save(str(savedialog[0]))
                        print(savedialog[0])
                    except FileNotFoundError:
                        self.statusBar().showMessage('알림 : 우편물 목록 저장되지 않음')

                    # Word 실행
                    os.system(f'start winword "{savedialog[0]}"')




    # 결재버튼 누를 시 할인그룹 우편리스트 추출하고 전표를 생성
    def fee_to_docx(self):
        # postinformation 데이터가 있는지 확인
        postdb = sqlite3.connect("postinformation.db")
        cursor = postdb.cursor()
        # postage 존재하는지 Execute
        try:
            cursor.execute('SELECT * FROM postage LIMIT 1')
        except:
            # 오류 발생 시
            msg_nodb = QMessageBox()
            msg_nodb.warning(self, "오류", "우체국 요금데이터가 없습니다. 갱신해 주세요.")
            output_nodb = msg_nodb.show()
            if output_nodb == QMessageBox.Ok:
                msg_nodb.close()
                self.mainmenu_statusBar.showMessage('리셋완료')
        else:
            postdb.close()
            # 데이터 긁어오기
            self.table_postlist_sel.selectAll()

            # 긁어온 데이터가 존재하지 않을 경우 확인
            try:
                if len(self.table_postlist_sel.selectionModel().selectedRows()) == 0:
                    raise Exception
            except Exception as e:  # 예외가 발생했을 때 실행됨
                msg_noentry = QMessageBox()  # 메세지박스 생성
                msg_noentry.warning(self, "오류", "결재할 우편물을 선택해 주세요")  # 메세지박스 설정
                output_noentry = msg_noentry.show()  # 메세지박스 실행
                if output_noentry == QMessageBox.Ok:  # OK 누르면 닫기
                    msg_noentry.close()
            else:
                # 긁어온 데이터가 있을 경우
                readypostlist = self.getselectedrowid_adv(self.table_postlist_sel)
                print('[res]', readypostlist)
                # 데이터프레임에 헤더값 넣기
                labels = []
                for c in range(self.table_postlist_sel.columnCount()):
                    headeritem = self.table_postlist_sel.horizontalHeaderItem(c).text()
                    labels.append(headeritem)
                # 데이터프레임 생성
                print('현재라벨 :', labels)
                postdf = pd.DataFrame(readypostlist)
                postdf.columns = labels

                if self.is_groupified == True:
                    #### SQLite DB 불러오기 ####
                    # 묶음할인대상 우편물은 grouppostlist.db 에 옮겨서 관리
                    if os.path.isfile("grouppostlist.db"):  # DB파일이 있을 경우
                        groupdb = sqlite3.connect("grouppostlist.db")
                        groupcursor = groupdb.cursor()
                    else:  # DB파일이 없을 경우
                        self.statusBar().showMessage('DB가 없으므로 DB파일을 생성합니다.')
                        groupdb = sqlite3.connect("grouppostlist.db")
                        groupcursor = groupdb.cursor()
                        groupcursor.execute(
                            "CREATE TABLE groupedlist(id integer PRIMARY KEY AUTOINCREMENT, 접수날짜 text, 사원번호 text, 접수자명 text, 부서명 text, 보내는사람 text, 받는사람 text, 주소 text, 우편번호 text, 제목 text, 수량 integer, 종류 text, 긴급여부 text, 가로길이 real, 세로길이 real, 우편중량 real, 전화번호 text, 높이 real, 메모 text, 할인타입 text, 할인타입그룹 integer, 결재여부 , 반려여부 boolean, 비용 text)")
                        groupdb.commit()

                    if '접수불가' in postdf['비용'].values:
                        # 오류 메세지박스 생성(단순객체) #
                        msg_forbid = QMessageBox()  # 메세지박스 생성
                        msg_forbid.warning(self, "오류", '접수불가한 우편물이 있습니다 우편물의 비용을 확인해주세요.')  # 메세지박스 경고옵션 설정
                        output_forbid = msg_forbid.show()  # 메세지박스 실행
                        if output_forbid == QMessageBox.Ok:  # OK 누르면 닫기
                            msg_forbid.close()
                    else:
                        # 할인타입그룹이 있는 레코드만 groupdf 에 저장
                        groupdf = postdf.loc[(postdf['할인타입그룹'] != 'nan')]
                        nan_groupdf = postdf.loc[(postdf['할인타입그룹'] == 'nan')]

                        # DF 를 SQlite 형태로 변환 후 APPEND
                        groupdf.to_sql('groupedlist', groupdb, if_exists='replace', index=False)
                        groupdb.commit()

                        # 결재테이블 내 모든 데이터 선택
                        self.table_postlist_sel.selectAll()
                        # 결재된 우편물의 레코드의 결재여부에 1 표시 및 DB 업데이트
                        residlist = self.getselectedrowid(self.table_postlist_sel)

                        # list 의 id 값을 기반으로 sqlite 에서 레코드 변경
                        for idvalue in residlist:
                            self.cursor.execute("UPDATE postlist SET 결재여부 = 1 WHERE id=?", (idvalue,))
                        # 커밋 수행
                        self.listdb.commit()

                        ###### 여기서부터 정리작업 시작 ######
                        # 테이블 업데이트
                        self.loadpostlist()
                        # 결재대상 테이블 지우기
                        self.table_postlist_sel.clear()

                        # 비용 칼럼의 합을 표시하기 위해 str 타입을 float 타입으로 변환
                        nan_groupdf.loc[:, '비용'] = nan_groupdf['비용'].astype(float)
                        # 전표 데이터프레임 생성 (as_index 를 false 로 해서
                        resdf = nan_groupdf.groupby(['종류', '우편중량'], as_index=False)['비용'].sum()
                        print('[resdf]\n', resdf)
                        # 총 비용 레코드 생성
                        total_cost = resdf['비용'].sum()
                        total_row = {'종류': '총 비용', '우편중량': '', '비용': total_cost}
                        # 레코드 데이터프레임에 추가
                        resdf = pd.concat([resdf, pd.DataFrame([total_row])], ignore_index=True)

                        resdf = pd.DataFrame(resdf)
                        print('[resdf]\n', resdf)

                        # 결재 단위 표 생성
                        confirmdata = [['승인', '?', '결재', 'A', 'B']]
                        confirmdf = pd.DataFrame(confirmdata, columns=['승인', '승인자', '결재', '결재자1', ' 결재자2'])
                        print('[confirmdf]\n', confirmdf)

                        # Word 문서 생성
                        doc = Document()

                        # 제목 추가하기
                        doc.add_heading('결재전표', 0)

                        # 내용 추가하기
                        description_1 = doc.add_paragraph("우편물의 결재요금을 다음과 같이 설명합니다.")

                        # 요금표 생성
                        restable = doc.add_table(resdf.shape[0] + 1, resdf.shape[1])
                        # 요금표 테두리 생성
                        restable.style = 'Table Grid'
                        # 표의 1번쨰 줄부터 레코드 추가
                        # 칼럼 추가
                        for i, colname in enumerate(resdf.columns):
                            restable.cell(0, i).text = colname
                        # 내용 추가
                        for i in range(resdf.shape[0]):
                            for j in range(resdf.shape[1]):
                                restable.cell(i + 1, j).text = str(resdf.values[i, j])
                        # 테이블의 글자크기를 7로 변경 (코드참조 : https://stackoverflow.com/questions/43007725/python-docx-how-to-change-table-font-size)
                        for row in restable.rows:  # 테이블의 각 줄
                            for cell in row.cells:  # 테이블 각 줄의 셀
                                paragraphs = cell.paragraphs  # 테이블 각 줄의 셀의 문단그룹
                                for paragraph in paragraphs:  # 테이블 각 줄의 셀의 문단
                                    for word in paragraph.runs:  # 문단의 모든 글자
                                        word.font.size = Pt(8)

                        # 마지막 텍스트 추가
                        text_to_add = datetime.today().strftime(
                            '%Y년 %m월 %d일 %H시 %M분 %S초') + "에 취합한 우편물의 발송에 필요한 금액을 결재하고자 하오니 검토후 재가하여 주시기 바랍니다."
                        description_last = doc.add_paragraph(text_to_add)
                        description_last.alignment = 1  # 가운데 정렬

                        # 결재테이블 추가
                        # 결재표 생성
                        confirmtable = doc.add_table(confirmdf.shape[0] + 1, confirmdf.shape[1])
                        # 결재표 테두리 생성
                        confirmtable.style = 'Table Grid'
                        # 표의 1번쨰 줄부터 레코드 추가
                        # 칼럼 추가
                        for i, colname in enumerate(confirmdf.columns):
                            confirmtable.cell(0, i).text = colname
                        #
                        for i in range(confirmdf.shape[0]):
                            for j in range(confirmdf.shape[1]):
                                confirmtable.cell(i + 1, j).text = str(confirmdf.values[i, j])
                        # 테이블의 글자크기를 7로 변경 (코드참조 : https://stackoverflow.com/questions/43007725/python-docx-how-to-change-table-font-size)
                        for row in confirmtable.rows:  # 테이블의 각 줄
                            for cell in row.cells:  # 테이블 각 줄의 셀
                                paragraphs = cell.paragraphs  # 테이블 각 줄의 셀의 문단그룹
                                for paragraph in paragraphs:  # 테이블 각 줄의 셀의 문단
                                    for word in paragraph.runs:  # 문단의 모든 글자
                                        word.font.size = Pt(8)

                        # Word 문서 저장
                        savename = datetime.today().strftime('%Y-%m-%d_%H-%M-%S_') + '결재전표.docx'
                        print('저장된 파일명 :', savename)

                        # 저장 다이얼로그 생성
                        savedialog = QFileDialog.getSaveFileName(self, '결재전표 문서 저장', savename, 'word 파일(*.docx)')

                        # 저장 실행
                        try:
                            doc.save(str(savedialog[0]))
                            print(savedialog[0])
                        except FileNotFoundError:
                            self.statusBar().showMessage('알림 : 결재전표 생성되지 않음')

                        # Word 실행
                        os.system(f'start winword "{savedialog[0]}"')

                else:
                    msg_forbid = QMessageBox()  # 메세지박스 생성
                    msg_forbid.warning(self, "실행제한", "우편물 비용 절약을 위해 [4. 묶음할인 설정] 까지 진행한 다음에 결재해주세요.")  # 메세지박스 설정
                    output_forbid = msg_forbid.show()  # 메세지박스 실행
                    if output_forbid == QMessageBox.Ok:  # OK 누르면 닫기
                        msg_forbid.close()




